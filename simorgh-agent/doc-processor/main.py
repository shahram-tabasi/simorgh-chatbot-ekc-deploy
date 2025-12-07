# main.py
# Universal Document Processor (PDF, Image, Word, Excel, Text → Markdown)
# FastAPI + EasyOCR + pdfplumber + camelot + pandas

import os
import time
import uuid
import warnings
import re
import asyncio
from datetime import datetime
from pathlib import Path

# Suppress warnings
warnings.filterwarnings('ignore')

# FastAPI imports
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware

# Document processing imports
import pdfplumber
import easyocr
from pdf2image import convert_from_path
import pandas as pd
from PIL import Image, ImageEnhance
import numpy as np
import aiofiles
from docx import Document

# ------------------------------------------------------------------
# FastAPI App Configuration
# ------------------------------------------------------------------
app = FastAPI(
    title="Universal Document → Markdown Converter",
    description="Convert PDF, Images, Word, Excel, and Text files to Markdown with Persian/English OCR support",
    version="1.2.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ------------------------------------------------------------------
# Helper Classes
# ------------------------------------------------------------------
class FileValidator:
    MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
    ALLOWED_EXTENSIONS = {
        '.pdf', '.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif',
        '.docx', '.doc', '.xlsx', '.xls', '.csv', '.txt', '.md'
    }

    @classmethod
    def validate(cls, file_path: Path) -> tuple[bool, str]:
        if not file_path.exists():
            return False, "File does not exist"
        if file_path.stat().st_size > cls.MAX_FILE_SIZE:
            return False, "File size exceeds 100MB"
        if file_path.stat().st_size == 0:
            return False, "File is empty"
        if file_path.suffix.lower() not in cls.ALLOWED_EXTENSIONS:
            return False, f"Format {file_path.suffix} not supported"
        return True, "Valid"


class PerformanceMonitor:
    def __init__(self):
        self.metrics = {
            'total_requests': 0,
            'successful': 0,
            'failed': 0,
            'total_time': 0.0,
            'by_type': {}
        }

    def record(self, success: bool, duration: float, file_type: str = None):
        self.metrics['total_requests'] += 1
        self.metrics['total_time'] += duration
        if success:
            self.metrics['successful'] += 1
        else:
            self.metrics['failed'] += 1
        if file_type:
            self.metrics['by_type'][file_type] = self.metrics['by_type'].get(file_type, 0) + 1

    def get_stats(self) -> dict:
        total = self.metrics['total_requests']
        return {
            **self.metrics,
            'avg_time': round(self.metrics['total_time'] / total, 2) if total > 0 else 0,
            'success_rate': round((self.metrics['successful'] / total * 100), 2) if total > 0 else 0
        }


monitor = PerformanceMonitor()


# ------------------------------------------------------------------
# Main Document Processor
# ------------------------------------------------------------------
class UniversalDocumentProcessor:
    def __init__(self, output_folder="processed_docs"):
        self.output_folder = Path(output_folder)
        self.output_folder.mkdir(exist_ok=True)
        self.ocr_reader = None
        print(f"Output directory: {self.output_folder.absolute()}")

    def get_ocr_reader(self):
        if self.ocr_reader is None:
            print("Loading EasyOCR (first time only, may take a while)...")
            self.ocr_reader = easyocr.Reader(['fa', 'en'], gpu=False, verbose=False)
            print("OCR Ready!")
        return self.ocr_reader

    def fix_text(self, text: str) -> str:
        """Clean and normalize text (Persian numbers to Latin, spacing fixes)"""
        if not text:
            return text
        # Convert Persian digits to Latin
        text = text.translate(str.maketrans('۰۱۲۳۴۵۶۷۸۹', '0123456789'))
        # Fix common spacing issues
        text = re.sub(r'([0-9]+)\s+([a-zA-Z]+)', r'\2 \1', text)
        text = re.sub(r'([0-9]+)\s*[kK]\s*[Vv]', r'\1kV', text)
        text = re.sub(r'([0-9]+)\s*[kK]\s*[Aa]', r'\1kA', text)
        text = re.sub(r'([0-9]+)\s*[Hh]z', r'\1Hz', text)
        return re.sub(r'\s+', ' ', text).strip()

    async def process_pdf(self, file_path: Path) -> str:
        """Process PDF using pdfplumber for text and tables"""
        parts = []
        with pdfplumber.open(file_path) as pdf:
            for num, page in enumerate(pdf.pages, 1):
                parts.append(f"## Page {num}\n\n")

                # Extract tables
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        cleaned = [[self.fix_text(str(c)) if c else "" for c in row] for row in table]
                        if len(cleaned) > 1:
                            header = cleaned[0]
                            parts.append("| " + " | ".join(header) + " |\n")
                            parts.append("| " + " | ".join(["---"] * len(header)) + " |\n")
                            for row in cleaned[1:]:
                                while len(row) < len(header):
                                    row.append("")
                                parts.append("| " + " | ".join(row) + " |\n")
                            parts.append("\n")

                # Extract text
                text = page.extract_text()
                if text:
                    lines = [self.fix_text(l.strip()) for l in text.split('\n') if l.strip()]
                    if lines:
                        parts.append("\n\n".join(lines) + "\n\n")

                parts.append("---\n\n")
        return ''.join(parts)

    async def process_image(self, file_path: Path) -> str:
        """Process image using EasyOCR"""
        reader = self.get_ocr_reader()
        image = Image.open(file_path)
        img_array = np.array(image)

        # Convert to grayscale and enhance contrast
        if len(img_array.shape) == 3:
            gray = np.dot(img_array[..., :3], [0.2989, 0.5870, 0.1140]).astype(np.uint8)
        else:
            gray = img_array
        pil_gray = Image.fromarray(gray)
        enhancer = ImageEnhance.Contrast(pil_gray)
        enhanced = enhancer.enhance(2.0)
        img_final = np.array(enhanced)

        results = reader.readtext(img_final, paragraph=False, detail=1)
        # Sort top-to-bottom, left-to-right
        results = sorted(results, key=lambda x: (x[0][0][1], x[0][0][0]))

        texts = [self.fix_text(r[1]) for r in results if len(r) >= 3 and r[2] > 0.2]
        return '\n\n'.join(texts) if texts else "*No text detected*"

    async def process_word(self, file_path: Path) -> str:
        """Process Word document"""
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        for table in doc.tables:
            parts.append("\n")
            for i, row in enumerate(table.rows):
                cells = [c.text.strip() for c in row.cells]
                parts.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
            parts.append("\n")
        return '\n\n'.join(parts)

    async def process_excel(self, file_path: Path) -> str:
        """Process Excel/CSV file"""
        parts = []
        if file_path.suffix.lower() == '.csv':
            df = pd.read_csv(file_path)
            parts.append(df.to_markdown(index=False))
        else:
            xl = pd.ExcelFile(file_path)
            for sheet in xl.sheet_names:
                df = pd.read_excel(xl, sheet_name=sheet)
                parts.append(f"## {sheet}\n\n{df.to_markdown(index=False)}\n\n")
        return '\n\n'.join(parts)

    async def process_text(self, file_path: Path) -> str:
        """Process plain text file"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            return await f.read()

    async def process(self, file_path: Path, user_id: str, original_name: str) -> dict:
        """Main processing function"""
        start = time.time()
        try:
            is_valid, msg = FileValidator.validate(file_path)
            if not is_valid:
                raise ValueError(msg)

            suffix = file_path.suffix.lower()
            if suffix == '.pdf':
                content = await self.process_pdf(file_path)
                doc_type = 'pdf'
            elif suffix in {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif'}:
                content = await self.process_image(file_path)
                doc_type = 'image'
            elif suffix in {'.docx', '.doc'}:
                content = await self.process_word(file_path)
                doc_type = 'word'
            elif suffix in {'.xlsx', '.xls', '.csv'}:
                content = await self.process_excel(file_path)
                doc_type = 'excel'
            elif suffix in {'.txt', '.md'}:
                content = await self.process_text(file_path)
                doc_type = 'text'
            else:
                raise ValueError(f"Format {suffix} not supported")

            header = f"# {original_name}\n\n"
            header += f"**Source:** {original_name}\n"
            header += f"**User:** {user_id}\n"
            header += f"**Processed:** {datetime.now():%Y-%m-%d %H:%M:%S}\n\n---\n\n"

            markdown = header + content

            output_name = f"{user_id}_{Path(original_name).stem}.md"
            output_path = self.output_folder / output_name
            async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
                await f.write(markdown)

            duration = time.time() - start
            monitor.record(True, duration, doc_type)

            print(f"✅ Processed: {original_name} ({duration:.2f}s)")

            return {
                "success": True,
                "user_id": user_id,
                "original_filename": original_name,
                "output_filename": output_name,
                "output_path": str(output_path),
                "doc_type": doc_type,
                "chars": len(markdown),
                "content": markdown,  # Return markdown content
                "duration": round(duration, 2),
                "message": "Processing successful"
            }

        except Exception as e:
            duration = time.time() - start
            monitor.record(False, duration)
            print(f"❌ Error processing {original_name}: {e}")
            return {
                "success": False,
                "user_id": user_id,
                "original_filename": original_name,
                "error": str(e),
                "duration": round(duration, 2),
                "message": f"Error: {e}"
            }


processor = UniversalDocumentProcessor()


# ------------------------------------------------------------------
# API Routes
# ------------------------------------------------------------------
@app.get("/")
async def root():
    return {
        "message": "Universal Document → Markdown API is active!",
        "docs": "/docs",
        "health": "/health"
    }


@app.get("/health")
async def health():
    return {
        "status": "healthy",
        "service": "doc-processor",
        "version": "1.2.0"
    }


@app.post("/upload")
async def upload_file(
    user_id: str = Form(..., description="User ID"),
    file: UploadFile = File(..., description="PDF, Image, Word, Excel, or Text file")
):
    """
    Upload and process document to Markdown
    Returns markdown content and saved file path
    """
    # Save temporary file
    suffix = Path(file.filename).suffix.lower()
    temp_path = Path("temp") / f"{uuid.uuid4()}{suffix}"
    temp_path.parent.mkdir(exist_ok=True)

    try:
        async with aiofiles.open(temp_path, 'wb') as out_file:
            while content := await file.read(1024 * 1024):  # 1MB chunks
                await out_file.write(content)

        result = await processor.process(temp_path, user_id, file.filename)

        if result["success"]:
            return JSONResponse(content=result)
        else:
            raise HTTPException(status_code=400, detail=result["message"])

    finally:
        # Clean up temp file
        if temp_path.exists():
            temp_path.unlink()


@app.post("/upload/download")
async def upload_and_download(
    user_id: str = Form(...),
    file: UploadFile = File(...)
):
    """Upload and get markdown file as download"""
    suffix = Path(file.filename).suffix.lower()
    temp_path = Path("temp") / f"{uuid.uuid4()}{suffix}"
    temp_path.parent.mkdir(exist_ok=True)

    try:
        async with aiofiles.open(temp_path, 'wb') as out_file:
            while content := await file.read(1024 * 1024):
                await out_file.write(content)

        result = await processor.process(temp_path, user_id, file.filename)

        if result["success"]:
            return FileResponse(
                path=result["output_path"],
                filename=result["output_filename"],
                media_type="text/markdown"
            )
        else:
            raise HTTPException(status_code=400, detail=result["message"])

    finally:
        if temp_path.exists():
            temp_path.unlink()


@app.get("/stats")
async def stats():
    """Get processing statistics"""
    return monitor.get_stats()


if __name__ == "__main__":
    import uvicorn
    print("\n" + "="*60)
    print("Starting Document Processor Service...")
    print("Local:  http://localhost:8000")
    print("Docs:   http://localhost:8000/docs")
    print("="*60 + "\n")
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="info")
